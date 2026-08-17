#!/usr/bin/env python3
"""Genera DOCX del reporte completo + resumen ejecutivo."""
import os
import re
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BASE = "/Users/david899/Documents/Default Project/proaco-evaluacion"

with open(os.path.join(BASE, "REPORTE_EVALUACION_PROACO.md"), encoding="utf-8") as f:
    reporte_completo = f.read()

doc = Document()
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(10.5)
style.paragraph_format.space_after = Pt(4)
style.paragraph_format.line_spacing = 1.15

# Colores semáforo
GREEN = RGBColor(0x15, 0x57, 0x24)
YELLOW = RGBColor(0x85, 0x64, 0x04)
RED = RGBColor(0x72, 0x1C, 0x24)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
HEADER_BG = RGBColor(0x0F, 0x4C, 0x81)
ROW_ALT = RGBColor(0xF0, 0xF4, 0xF9)

def set_cell_shading(cell, color):
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), f'{color[0]:02X}{color[1]:02X}{color[2]:02X}')
    shading.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(shading)

def add_heading_styled(doc, text, level):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = HEADER_BG
    return h

def score_color(score_str):
    try:
        v = float(score_str)
    except:
        return None
    if v >= 0.75:
        return GREEN
    elif v >= 0.50:
        return YELLOW
    return RED

def make_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Header
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.bold = True
                run.font.color.rgb = WHITE
                run.font.size = Pt(8)
        set_cell_shading(cell, HEADER_BG)
    # Data
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(val)
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER if c_idx > 0 else WD_ALIGN_PARAGRAPH.LEFT
                for run in p.runs:
                    run.font.size = Pt(8)
            # Color por score en columnas numéricas
            if c_idx > 0:
                color = score_color(val)
                if color:
                    for run in cell.paragraphs[0].runs:
                        run.font.color.rgb = color
                        run.bold = True
            # Filas alternadas
            if r_idx % 2 == 1:
                for c in range(len(headers)):
                    set_cell_shading(table.rows[r_idx + 1].cells[c], ROW_ALT)
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)
    return table

# --- Parsear el markdown completo para extraer tablas ---
# Simplificado: usamos el reporte completo + resumen ya formateado
doc.add_heading('Reporte de Evaluación Voicebot Grupo Proaco — Inbound + Outbound', 0)
p = doc.add_paragraph()
run = p.add_run('Fecha: 2026-08-12 | Flows: inbound (cliente llama) y outbound (bot llama a leads) | Juez: Qwen 2.5 7B / 7B-Coder (Ollama local)')
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

# Leyenda
doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run('Leyenda semáforo: ')
run.bold = True
run.font.size = Pt(9)
for label, color in [('🟢 Bueno (≥0.75)', GREEN), ('🟡 Regular (0.50–0.74)', YELLOW), ('🔴 Crítico (<0.50)', RED)]:
    run = p.add_run('  ' + label + '  ')
    run.font.size = Pt(9)
    run.font.color.rgb = color

doc.add_page_break()

# ===== PARTE 1: REPORTE COMPLETO (resumido con tablas) =====
add_heading_styled(doc, '1. Resumen Ejecutivo', 1)

doc.add_paragraph('Se evaluaron 13 llamadas inbound reales + 6 touchpoints outbound (QA) contra reglas de negocio Proaco.')
doc.add_paragraph('Resultado general: Inbound 77% promedio (reglas Proaco, juez qwen2.5:7b). Outbound 0% en métricas de flow saliente (no hay llamadas outbound reales).')

add_heading_styled(doc, '2. Resultados Inbound — Suite 3 Reglas Proaco (juez qwen2.5:7b)', 2)
make_table(doc,
    ['Métrica', 'Promedio', 'Estado', 'Lectura'],
    [
        ['Listado máx 3', '0.93', '🟢', 'Respeta límite 3 items'],
        ['Tono español rioplatense', '0.82', '🟢', 'Amable, claro, rioplatense'],
        ['Deriva a la web', '0.83', '🟢', 'Deriva cuando no sabe'],
        ['Detección de intent', '0.75', '🟡', 'Detecta flujo, algún desvío'],
        ['Despedida oficial', '0.75', '🟡', 'Falta en 25% llamadas'],
        ['Saludo oficial', '0.62', '🟡', 'Ausente en 23% llamadas'],
        ['Pedido de contacto (1 vez)', '0.59', '🔴', 'Se repite o no pide'],
    ],
    [3, 2, 2, 7]
)

add_heading_styled(doc, '3. Comparación Juez1 (qwen2.5:7b) vs Juez2 (qwen2.5-coder:7b)', 2)
make_table(doc,
    ['Métrica', 'Juez1 (qwen2.5:7b)', 'Juez2 (coder)', 'Δ'],
    [
        ['Saludo oficial', '0.62', '0.35', '-0.27'],
        ['Despedida oficial', '0.75', '0.49', '-0.26'],
        ['Detección intent', '0.75', '0.77', '+0.02'],
        ['Pedido contacto 1 vez', '0.59', '0.65', '+0.06'],
        ['Listado máx 3', '0.93', '0.54', '-0.39'],
        ['Tono español', '0.82', '0.87', '+0.05'],
        ['Deriva web', '0.83', '0.89', '+0.06'],
    ],
    [4, 3, 3, 2]
)

add_heading_styled(doc, '4. Heurísticas Deterministas (sin LLM)', 2)
make_table(doc,
    ['Métrica', 'Promedio', 'Estado'],
    [
        ['Saludo presente', '0.31', '🔴'],
        ['Despedida presente', '0.85', '🟢'],
        ['URL derivación', '0.31', '🔴'],
        ['Sin datos AySA', '0.92', '🟢'],
        ['Sin loop transfer', '0.92', '🟢'],
        ['Tono respetuoso', '1.00', '🟢'],
        ['Adherencia español', '1.00', '🟢'],
    ],
    [7, 2, 3]
)

add_heading_styled(doc, '5. LLM-Judges Adicionales (Opik genéricos)', 2)
make_table(doc,
    ['Métrica', 'Promedio', 'Estado'],
    [
        ['Moderation', '0.00', '🔴'],
        ['Frustración usuario', '0.09', '🔴'],
        ['Correctitud tools', '0.30', '🔴'],
        ['Utilidad', '0.32', '🔴'],
        ['Relevancia respuesta', '0.35', '🔴'],
        ['Completitud tarea', '0.52', '🔴'],
        ['Coherencia conversacional', '0.64', '🟡'],
        ['Completitud sesión', '0.60', '🟡'],
    ],
    [5, 2, 3]
)

add_heading_styled(doc, '6. Outbound — Heurísticas + LLM-Judges', 2)
doc.add_paragraph('Nota: Las 6 llamadas son touchpoints de QA inbound; no hay llamadas outbound reales. Métricas de flow saliente = 0%.')
make_table(doc,
    ['Métrica (heurística)', 'Promedio', 'Estado', 'Nota'],
    [
        ['Se presenta como Proaco', '1.00', '🟢', 'Match "Proaco" en saludo inbound'],
        ['Menciona motivo', '0.17', '🔴', 'Solo 1/6 menciona campaña'],
        ['Pide consentimiento', '0.00', '🔴', 'Flow inbound no lo implementa'],
        ['Maneja no-interés', '1.00', '🟢', 'Falso positivo por "gracias"'],
        ['Ofrece agendar cita', '0.00', '🔴', 'Flow inbound no agenda'],
        ['Listado máx 3', '0.83', '🟢', '1 llamada lista >3'],
        ['Tono respetuoso', '1.00', '🟢', 'Sin problemas'],
    ],
    [4, 2, 2, 6]
)

add_heading_styled(doc, '7. LLM-Judges Outbound (GEval)', 2)
make_table(doc,
    ['Métrica', 'Promedio', 'Estado'],
    [
        ['Cliente no interesado', '0.13', '🔴'],
        ['Agendamiento cita', '0.00', '🔴'],
    ],
    [3, 2, 2]
)

doc.add_page_break()

# ===== RESUMEN EJECUTIVO PARA EQUIPO =====
add_heading_styled(doc, 'Resumen Ejecutivo para el Equipo', 1)

doc.add_paragraph('Evaluación Voicebot Proaco — Agosto 2026')
doc.add_paragraph('Qué se evaluó: 13 llamadas inbound reales + 6 touchpoints outbound (QA) contra reglas de negocio Proaco.')
doc.add_paragraph('Resultado general: Inbound 77% promedio (reglas Proaco, juez qwen2.5:7b). Outbound 0% en métricas de flow saliente (no hay llamadas outbound reales).')

add_heading_styled(doc, '✅ Lo que funciona bien (🟢 ≥ 0.75)', 2)
make_table(doc,
    ['Métrica', 'Score', 'Estado'],
    [
        ['Listado máx 3 (juez1)', '0.93', '🟢'],
        ['Tono español (juez1)', '0.82', '🟢'],
        ['Deriva web (juez1)', '0.83', '🟢'],
        ['Detección intent (juez1)', '0.75', '🟢'],
        ['Tono español (juez2)', '0.87', '🟢'],
        ['Deriva web (juez2)', '0.89', '🟢'],
    ],
    [6, 2, 2]
)

add_heading_styled(doc, '⚠️ Qué hay que arreglar (🟡 0.50–0.74 / 🔴 < 0.50)', 2)
make_table(doc,
    ['Métrica', 'Score', 'Estado', 'Acción'],
    [
        ['Saludo oficial (juez1)', '0.62', '🟡', 'Prompt obligatorio + test'],
        ['Pedido contacto (juez1)', '0.59', '🔴', 'Exactly-once al cierre + test'],
        ['Despedida oficial (juez1)', '0.75', '🟡', 'Fallback obligatorio'],
        ['Saludo oficial (juez2)', '0.35', '🔴', 'Revisar criterio juez2'],
        ['Despedida oficial (juez2)', '0.49', '🔴', 'Revisar criterio juez2'],
        ['Listado máx 3 (juez2)', '0.54', '🟡', 'Revisar criterio juez2'],
        ['Heurística maneja_no_interes', 'Bug', '🔴', 'Reescribir (match "gracias" saludo)'],
    ],
    [4, 2, 2, 7]
)

add_heading_styled(doc, '🚫 Outbound — Sin datos reales', 2)
doc.add_paragraph('Las 6 llamadas son inbound de QA, no outbound. Métricas flow saliente (presentación, consentimiento, agenda cita) = 0%.')
doc.add_paragraph('Próximo paso: configurar campaña outbound real en Lula con leads de prueba.')

add_heading_styled(doc, 'Roadmap', 2)
make_table(doc,
    ['Sprint', 'Foco', 'Entregable', 'Tiempo'],
    [
        ['1', 'Inbound crítico', 'Fix saludo, contacto, despedida, heurística → >0.80 reglas', '1-2 sem'],
        ['2', 'Outbound real', 'Campaña Lula leads → exportar → evaluar heurísticas + GEval', '2-3 sem'],
        ['3', 'Juez cloud + CI', 'Groq llama-3.3-70b (gratis) + pipeline automático', '1 sem'],
    ],
    [2, 4, 5, 2]
)

# Guardar
output_path = os.path.join(BASE, "REPORTE_EVALUACION_PROACO.docx")
doc.save(output_path)
print(f"✅ DOCX generado: {output_path}")